import os
import tempfile
import pathlib

import pytest

from docs_parser import docx_to_markdown_full


def create_sample_docx_with_image(docx_path, img_path):
    # create a simple docx with one paragraph and one picture
    from docx import Document

    doc = Document()
    doc.add_paragraph("Hello from test doc")
    doc.add_picture(img_path)
    doc.save(docx_path)


def create_sample_image(img_path):
    from PIL import Image

    img = Image.new("RGB", (10, 10), color=(255, 0, 0))
    img.save(img_path, format="PNG")


def test_docx_to_markdown_creates_md_and_images():
    with tempfile.TemporaryDirectory() as tmp:
        tmpdir = pathlib.Path(tmp)
        docx_path = tmpdir / "sample.docx"
        img_path = tmpdir / "img.png"

        create_sample_image(str(img_path))
        create_sample_docx_with_image(str(docx_path), str(img_path))

        out_md = tmpdir / "out.md"
        img_out_dir = tmpdir / "sample_images"

        docx_to_markdown_full(str(docx_path), str(out_md), str(img_out_dir))

        assert out_md.exists(), "Markdown file should be created"
        text = out_md.read_text(encoding="utf-8")
        assert "Hello from test doc" in text

        # image dir should contain at least one file
        files = list(img_out_dir.glob("*"))
        assert len(files) >= 1
        # ensure the file has an expected image extension
        assert any(p.suffix.lower() in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg', '.bin'} for p in files)
