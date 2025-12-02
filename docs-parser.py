from docx import Document
import os

def docx_to_markdown_full(docx_path, md_path, image_dir="images"):
    doc = Document(docx_path)
    md_lines = []

    # 이미지 저장 폴더 생성
    os.makedirs(image_dir, exist_ok=True)
    image_count = 1

    # 문단 처리 (제목 포함)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Heading → Markdown 제목 변환
        if para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", ""))
            except ValueError:
                level = 1
            md_lines.append("#" * level + " " + text)
        else:
            md_lines.append(text)

        # 하이퍼링크 처리
        for run in para.runs:
            if run.hyperlink:
                url = run.hyperlink.target
                link_text = run.text.strip() or url
                md_lines.append(f"[{link_text}]({url})")

    # 테이블 처리
    for t_idx, table in enumerate(doc.tables, start=1):
        md_lines.append(f"\n### Table {t_idx}\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("\n")

    # 이미지 추출
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image_filename = f"{image_dir}/image_{image_count}.png"
            with open(image_filename, "wb") as f:
                f.write(image_data)
            md_lines.append(f"![image_{image_count}]({image_filename})")
            image_count += 1

    # Markdown 저장
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(md_lines))

# 실행 예시
docx_to_markdown_full(
    "/home/dnshine/python-files/API-TEST.docx",
    "/home/dnshine/python-files/output.md",
    "/home/dnshine/python-files/images"
)