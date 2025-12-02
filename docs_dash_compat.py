from docx import Document
import os
import argparse
import logging
import pathlib
import sys


def docx_to_markdown_full(docx_path, md_path, image_dir="images"):
    """Convert a single .docx file to Markdown.

    - docx_path: path to source .docx
    - md_path: path to write resulting markdown (.md)
    - image_dir: path to store any images (will be created)
    """
    doc = Document(docx_path)
    md_lines = []

    # 이미지 저장 폴더 생성
    os.makedirs(image_dir, exist_ok=True)
    image_count = 1

    # 문단 처리 (제목 포함)
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # Heading → Markdown 제목 변환
        try:
            style_name = para.style.name
        except Exception:
            style_name = ""

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.replace("Heading ", ""))
            except ValueError:
                level = 1
            md_lines.append("#" * level + " " + text)
        else:
            md_lines.append(text)

        # 하이퍼링크 처리 (간단히 처리)
        for run in para.runs:
            if hasattr(run, "hyperlink") and run.hyperlink:
                try:
                    url = run.hyperlink.target
                    link_text = run.text.strip() or url
                    md_lines.append(f"[{link_text}]({url})")
                except Exception:
                    # best-effort — ignore malformed hyperlink
                    pass

    # 테이블 처리
    for t_idx, table in enumerate(doc.tables, start=1):
        md_lines.append(f"\n### Table {t_idx}\n")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |")
        md_lines.append("\n")

    # 이미지 추출 (use relationship blobs) — detect extensions and avoid overwrites
    def _unique_filename(dirpath, base, ext):
        # ensure directory exists
        os.makedirs(dirpath, exist_ok=True)
        candidate = f"{base}{ext}"
        i = 1
        while os.path.exists(os.path.join(dirpath, candidate)):
            candidate = f"{base}_{i}{ext}"
            i += 1
        return candidate

    for rel in getattr(doc.part, "rels", {}).values():
        try:
            if "image" in rel.reltype:
                image_data = rel.target_part.blob

                # try to get extension from partname (eg: /word/media/image1.png)
                ext = None
                try:
                    partname = getattr(rel.target_part, 'partname', None)
                    if partname:
                        ext = pathlib.Path(partname).suffix
                except Exception:
                    ext = None

                # fallback to content_type if no ext
                if not ext:
                    try:
                        ctype = getattr(rel.target_part, 'content_type', '')
                        if '/' in ctype:
                            subtype = ctype.split('/')[1]
                            # handle image/svg+xml
                            subtype = subtype.split('+')[0]
                            ext = '.' + ( 'jpg' if subtype == 'jpeg' else subtype )
                    except Exception:
                        ext = '.bin'

                if not ext:
                    ext = '.bin'

                # pick a base name that is more descriptive than just image_ number
                base = f"image_{image_count}"
                fname = _unique_filename(image_dir, base, ext)
                image_filename = os.path.join(image_dir, fname)
                with open(image_filename, "wb") as f:
                    f.write(image_data)
                # add relative path to markdown (make path relative to md file)
                relpath = os.path.relpath(image_filename, os.path.dirname(md_path))
                md_lines.append(f"![{base}]({relpath})")
                image_count += 1
        except Exception:
            # ignore image extraction errors for robustness
            continue

    # Markdown 저장
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n\n".join(md_lines))


def process_directory(input_dir, output_dir, image_subdir_name="images", recursive=False, logger=None):
    """Process all .docx files in input_dir and write .md files into output_dir.

    For each file Lorem.docx, this will create output_dir/Lorem.md and images at
    output_dir/Lorem_images/ (or the provided image_subdir_name).
    """
    p = pathlib.Path(input_dir)
    if not p.exists():
        raise FileNotFoundError(f"Input directory not found: {input_dir} — please check the path")
    if not p.is_dir():
        raise NotADirectoryError(f"Input path exists but is not a directory: {input_dir}")

    out_p = pathlib.Path(output_dir)
    try:
        out_p.mkdir(parents=True, exist_ok=True)
    except PermissionError:
        raise PermissionError(f"Cannot create or write to output directory: {output_dir} — check permissions")
    except Exception as e:
        raise OSError(f"Failed to create output directory {output_dir}: {e}")

    pattern = "**/*.docx" if recursive else "*.docx"
    files = list(p.glob(pattern))
    processed = []

    for f in files:
        if not f.is_file():
            continue

        # construct output names
        stem = f.stem
        md_name = stem + ".md"
        md_path = out_p.joinpath(md_name)

        # image dir: per-file subdir under output_dir
        image_dir = out_p.joinpath(f"{stem}_{image_subdir_name}")
        try:
            docx_to_markdown_full(str(f), str(md_path), str(image_dir))
            processed.append((str(f), str(md_path), str(image_dir)))
            if logger:
                logger.info("Converted: %s -> %s (images: %s)", f, md_path, image_dir)
            else:
                print(f"Converted: {f} -> {md_path} (images: {image_dir})")
        except Exception as e:
            if logger:
                logger.exception("Failed to convert %s", f)
            else:
                print(f"Failed to convert {f}: {e}", file=sys.stderr)

    if not processed and logger:
        logger.warning("No .docx files were found in %s (pattern=%s)", input_dir, pattern)

    return processed


def __main__():
    ap = argparse.ArgumentParser(description="Convert .docx files into Markdown files (single-file or directory batch mode)")

    group = ap.add_mutually_exclusive_group(required=True)
    group.add_argument("--input-dir", help="Directory with .docx files to convert")
    group.add_argument("--file", help="Single .docx file to convert")

    ap.add_argument("--output-dir", default=None, help="Directory to write .md files and images (for --file, defaults to parent folder)" )
    ap.add_argument("--images-subdir", default="images", help="Name for per-file images subdirectory suffix (default 'images')")
    ap.add_argument("--recursive", action="store_true", help="Recurse into subdirectories to find .docx files")
    ap.add_argument("--quiet", action="store_true", help="Minimal output")
    ap.add_argument("--verbose", action="store_true", help="Show detailed processing info (INFO level)")
    args = ap.parse_args()

    # configure logging
    log_level = logging.WARNING
    if args.quiet:
        log_level = logging.ERROR
    elif args.verbose:
        log_level = logging.INFO
    logging.basicConfig(level=log_level, format="[%(levelname)s] %(message)s")
    logger = logging.getLogger("docs-parser")

    try:
        if args.file:
            # single-file mode
            fpath = pathlib.Path(args.file)
            if not fpath.exists() or not fpath.is_file():
                logger.error("Input file does not exist or is not a file: %s", args.file)
                sys.exit(2)

            # determine output directory
            out_dir = args.output_dir or str(fpath.parent)
            pathlib.Path(out_dir).mkdir(parents=True, exist_ok=True)

            stem = fpath.stem
            md_path = pathlib.Path(out_dir).joinpath(stem + ".md")
            image_dir = pathlib.Path(out_dir).joinpath(f"{stem}_{args.images_subdir}")
            try:
                docx_to_markdown_full(str(fpath), str(md_path), str(image_dir))
                logger.info("Converted file: %s -> %s (images: %s)", fpath, md_path, image_dir)
            except Exception:
                logger.exception("Failed to convert %s", fpath)
                sys.exit(1)
        else:
            # directory/batch mode
            if not args.input_dir:
                logger.error("--input-dir must be provided in directory mode.")
                sys.exit(2)
            results = process_directory(args.input_dir, args.output_dir or '.', image_subdir_name=args.images_subdir, recursive=args.recursive, logger=logger)
            if not args.quiet:
                logger.info("Processed %d files.", len(results))
    except FileNotFoundError as e:
        logger.error(str(e))
        logger.info("Verify the path and try again.")
        sys.exit(2)
    except NotADirectoryError as e:
        logger.error(str(e))
        sys.exit(2)
    except PermissionError as e:
        logger.error(str(e))
        sys.exit(3)
    except OSError as e:
        logger.error(str(e))
        sys.exit(3)
    except Exception as e:
        logger.exception("Unexpected error: %s", e)
        sys.exit(1)

if __name__ == '__main__':
    __main__()
